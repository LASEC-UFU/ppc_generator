import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Define bordas para uma célula"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Cria o elemento de bordas
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), edge_data.get('color', '2E5090'))
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def set_cell_background(cell, color):
    """Define cor de fundo da célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def parse_latex_table(latex_content):
    """Extrai dados da tabela LaTeX"""
    rows = []
    current_period = ""
    
    # Remove comentários e linhas vazias
    lines = [line.strip() for line in latex_content.split('\n') 
             if line.strip() and not line.strip().startswith('%')]
    
    # Processa linha por linha
    for line in lines:
        # Ignora linhas de configuração e comandos (mas NÃO SetCell de dados)
        if any(cmd in line for cmd in ['\\begin{', '\\end{', '\\hline', 
                                         'colspec', 'rowsep', 'note{', 'caption', 
                                         'label', 'theme', 'rowhead', 'cells  =', 
                                         'row{', 'vlines', 'hlines']):
            continue
        
        # Verifica se é uma linha com SetCell[r=X] que define o período
        if '\\SetCell[r=' in line and 'textordmasculine' in line:
            # Extrai o período da linha SetCell
            match = re.search(r'(\d+)\\textordmasculine', line)
            if match:
                current_period = match.group(1) + 'º'
            continue
            
        # Detecta linhas de dados (contém &)
        if '&' in line and not line.startswith('\\SetRow'):
            # Remove \\ final e divide por &
            line = line.replace('\\\\', '').strip()
            cells = [cell.strip() for cell in line.split('&')]
            
            if len(cells) >= 10:  # Linha válida com dados completos
                # Limpa formatação LaTeX básica
                cleaned_cells = []
                for cell in cells:
                    # Remove comandos LaTeX comuns
                    cell = re.sub(r'\\textbf\{([^}]+)\}', r'\1', cell)
                    cell = re.sub(r'\\textit\{([^}]+)\}', r'\1', cell)
                    cell = re.sub(r'\\textordmasculine\{\}', 'º', cell)
                    cell = re.sub(r'\\TblrNote\{[^}]+\}', '', cell)
                    cell = re.sub(r'\$\\times\$', '×', cell)
                    cell = re.sub(r'\$\\dagger [0-9]\$', '', cell)
                    cell = re.sub(r'\$\\bullet\$', '•', cell)
                    cell = cell.replace('\\textit{', '').replace('}', '')
                    cell = cell.strip()
                    cleaned_cells.append(cell)
                
                # Se a primeira célula estiver vazia, usa o período atual
                if not cleaned_cells[0] or cleaned_cells[0] == '':
                    cleaned_cells[0] = current_period
                else:
                    # Atualiza o período atual se a célula não estiver vazia
                    if cleaned_cells[0] not in ['', 'Disciplinas Optativas', 'Atividade Acadêmicas Complementares', 'Atividade de Conclusão de Curso']:
                        current_period = cleaned_cells[0]
                
                rows.append(cleaned_cells)
    
    return rows

def create_word_table(data, output_file='tabela_fluxo_curricular_conforme_guia.docx'):
    """Cria documento Word com a tabela"""
    doc = Document()
    
    # Configurações da página
    section = doc.sections[0]
    section.page_width = Inches(11.69)  # A4 landscape
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Adiciona título
    title = doc.add_heading('Fluxo Curricular', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Cria tabela
    table = doc.add_table(rows=len(data) + 2, cols=11)
    table.style = 'Table Grid'
    
    # Cabeçalho principal (primeira linha)
    headers_main = ['PER', 'Componente Curricular', 'Natureza', 'Carga Horária', 
                    '', '', '', '', 'PREQ', 'CREQ', 'UA Oferta']
    
    for i, header in enumerate(headers_main):
        cell = table.rows[0].cells[i]
        cell.text = header
        # Mescla células para "Carga Horária"
        if i == 3:
            cell = table.rows[0].cells[3]
            cell.merge(table.rows[0].cells[7])
            cell.text = 'Carga Horária'
        
        # Formata cabeçalho
        set_cell_background(cell, '2E5090')
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # Subcabeçalho (segunda linha)
    subheaders = ['', '', '', 'CHT', 'CHP', 'CHD', 'CHE', 'TOT', '', '', '']
    for i, subheader in enumerate(subheaders):
        cell = table.rows[1].cells[i]
        if subheader:
            cell.text = subheader
            set_cell_background(cell, '2E5090')
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(7)
    
    # Preenche dados
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < 11:
                cell = table.rows[row_idx + 2].cells[col_idx]
                cell.text = str(cell_data)
                
                # Formata célula
                paragraph = cell.paragraphs[0]
                if col_idx in [0, 2, 3, 4, 5, 6, 7, 9, 10]:  # Células centralizadas
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.font.size = Pt(7)
                
                # Cor de fundo alternada
                if row_idx % 2 == 0:
                    set_cell_background(cell, 'E8E8E8')
    
    # Adiciona notas de rodapé
    doc.add_paragraph()
    notes = [
        "† 1: Os discentes deverão integralizar 345 horas de atividades extensionistas ao longo do curso.",
        "† 2: O ENADE é componente curricular obrigatório (Lei nº 10.861/2004).",
        "† 3: O discente deverá cursar 90 horas de atividades acadêmicas complementares.",
        "† 4: Para iniciar a ACC (300h), o discente deve ter integralizado 1875 horas em disciplinas obrigatórias.",
        "† 5: O discente deverá cursar, no mínimo, 90 horas em disciplinas optativas."
    ]
    
    for note in notes:
        p = doc.add_paragraph(note)
        p.style = 'List Bullet'
        run = p.runs[0]
        run.font.size = Pt(8)
    
    # Salva documento
    doc.save(output_file)
    print(f"Documento salvo como: {output_file}")

def main():
    # Lê arquivo LaTeX
    input_file = 'include/auto/tab_fluxo_curricular_conforme_guia.tex'  # Altere para o nome do seu arquivo
    
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            latex_content = f.read()
        
        # Extrai dados da tabela
        print("Processando tabela LaTeX...")
        data = parse_latex_table(latex_content)
        
        print(f"Encontradas {len(data)} linhas de dados")
        
        # Cria documento Word
        print("Criando documento Word...")
        create_word_table(data)
        
        print("Conversão concluída com sucesso!")
        
    except FileNotFoundError:
        print(f"Erro: Arquivo '{input_file}' não encontrado.")
        print("Certifique-se de que o arquivo está no mesmo diretório do script.")
    except Exception as e:
        print(f"Erro durante a conversão: {str(e)}")

if __name__ == "__main__":
    main()